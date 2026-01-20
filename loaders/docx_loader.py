from pathlib import Path
from docx import Document

def load_docx_file(file_path: str) -> list:
    """
    Extract text from Word (.docx) files.
    Returns list of document chunks with metadata.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"{file_path} not found")

    try:
        doc = Document(path)
    except Exception as e:
        raise ValueError(f"Error reading DOCX file {file_path}: {e}")

    docs = []
    content_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_data))
        if table_text:
            content_parts.append("TABLE:\n" + "\n".join(table_text))

    # Combine all content
    if content_parts:
        full_content = "\n\n".join(content_parts)
        docs.append({
            "content": full_content,
            "metadata": {
                "source": path.name,
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
        })

    return docs
